from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path

import structlog
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from maskflow.core.engine import MaskingEngine
from maskflow.core.types import AnalysisResult

logger = structlog.get_logger(__name__)


@dataclass(slots=True)
class _DocxStats:
    matches_found: int = 0
    matches_applied: int = 0
    matches_skipped: int = 0
    detector_counts: Counter[str] = field(default_factory=Counter)
    detector_timings_ms: Counter[str] = field(default_factory=Counter)


class DocxProcessor:
    def __init__(self, engine: MaskingEngine) -> None:
        self.engine = engine

    def analyze(self, source: Path | str) -> AnalysisResult:
        document = Document(str(source))

        total_matches_found = 0
        total_matches_applied = 0
        total_matches_skipped = 0
        detector_counts: Counter[str] = Counter()
        detector_timings_ms: Counter[str] = Counter()

        for text in self._iter_paragraph_texts(document):
            analysis = self.engine.analyze_text(text)

            total_matches_found += analysis.matches_found
            total_matches_applied += analysis.matches_applied
            total_matches_skipped += analysis.matches_skipped
            detector_counts.update(analysis.detector_counts)
            detector_timings_ms.update(analysis.detector_timings_ms)

        return AnalysisResult(
            matches_found=total_matches_found,
            matches_applied=total_matches_applied,
            matches_skipped=total_matches_skipped,
            detector_counts=dict(detector_counts),
            detector_timings_ms=dict(detector_timings_ms),
        )

    def _iter_paragraph_texts(self, document: DocxDocument) -> list[str]:
        values: list[str] = []

        for paragraph in document.paragraphs:
            if paragraph.text:
                values.append(paragraph.text)

        for table in document.tables:
            values.extend(self._iter_table_paragraph_texts(table))

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                if paragraph.text:
                    values.append(paragraph.text)
            for paragraph in section.footer.paragraphs:
                if paragraph.text:
                    values.append(paragraph.text)
            for table in section.header.tables:
                values.extend(self._iter_table_paragraph_texts(table))
            for table in section.footer.tables:
                values.extend(self._iter_table_paragraph_texts(table))

        return values

    def _iter_table_paragraph_texts(self, table: Table) -> list[str]:
        values: list[str] = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        values.append(paragraph.text)
                for nested_table in cell.tables:
                    values.extend(self._iter_table_paragraph_texts(nested_table))
        return values

    def process(self, source: Path | str, destination: Path | str) -> None:
        document = Document(str(source))
        self._process_document(document, stats=None)
        document.save(str(destination))

    def process_with_stats(
        self,
        source: Path | str,
        destination: Path | str,
    ) -> AnalysisResult:
        """Single-pass: маскирует и возвращает статистику за один проход.

        Устраняет двойное чтение файла по сравнению с analyze() + process().
        """
        document = Document(str(source))
        stats = _DocxStats()
        self._process_document(document, stats=stats)
        document.save(str(destination))

        return AnalysisResult(
            matches_found=stats.matches_found,
            matches_applied=stats.matches_applied,
            matches_skipped=stats.matches_skipped,
            detector_counts=dict(stats.detector_counts),
            detector_timings_ms=dict(stats.detector_timings_ms),
        )

    def _process_document(
        self,
        document: DocxDocument,
        stats: "_DocxStats | None",
    ) -> None:
        for paragraph in document.paragraphs:
            self._process_paragraph(paragraph, stats)

        for table in document.tables:
            self._process_table(table, stats)

        for section in document.sections:
            self._process_header_footer_paragraphs(section.header.paragraphs, stats)
            self._process_header_footer_paragraphs(section.footer.paragraphs, stats)

            for table in section.header.tables:
                self._process_table(table, stats)
            for table in section.footer.tables:
                self._process_table(table, stats)

    def _process_header_footer_paragraphs(
        self,
        paragraphs: list[Paragraph],
        stats: "_DocxStats | None",
    ) -> None:
        for paragraph in paragraphs:
            self._process_paragraph(paragraph, stats)

    def _process_table(
        self,
        table: Table,
        stats: "_DocxStats | None",
    ) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self._process_paragraph(paragraph, stats)
                for nested_table in cell.tables:
                    self._process_table(nested_table, stats)

    def _process_paragraph(
        self,
        paragraph: Paragraph,
        stats: "_DocxStats | None",
    ) -> None:
        """Маскируем абзац с сохранением форматирования runs по возможности.

        Стратегия (двухфазная):

        Фаза 1 — per-run masking.
            Каждый run маскируется отдельно. Форматирование (bold, italic,
            font, color и т. д.) сохраняется полностью. Применяется, когда
            чувствительное значение полностью умещается в одном run.

        Фаза 2 — merge fallback.
            Если склейка per-run результатов отличается от маскирования
            всего paragraph.text (т. е. чувствительное значение «разорвано»
            между runs), форматирование первого run расширяется на весь абзац,
            остальные runs зануляются. Потеря форматирования логируется.
        """
        if not paragraph.runs:
            return

        original = paragraph.text
        if not original:
            return

        # Фаза 1: per-run masking
        if stats is not None:
            run_pairs = [
                self.engine.process_with_stats(run.text) if run.text else (run.text, None)
                for run in paragraph.runs
            ]
            run_results = [masked for masked, _ in run_pairs]
            # Аккумулируем статистику по всем runs
            for _, analysis in run_pairs:
                if analysis is not None:
                    stats.matches_found += analysis.matches_found
                    stats.matches_applied += analysis.matches_applied
                    stats.matches_skipped += analysis.matches_skipped
                    stats.detector_counts.update(analysis.detector_counts)
                    stats.detector_timings_ms.update(analysis.detector_timings_ms)
            full_masked, full_analysis = self.engine.process_with_stats(original)
        else:
            run_results = [
                self.engine.process_text(run.text) if run.text else run.text
                for run in paragraph.runs
            ]
            full_masked = self.engine.process_text(original)

        per_run_joined = "".join(r or "" for r in run_results)

        if full_masked == original:
            # Нет чувствительных данных — ничего не делаем
            return

        if per_run_joined == full_masked:
            # Фаза 1 достаточна — применяем с сохранением форматирования
            for run, new_text in zip(paragraph.runs, run_results, strict=True):
                run.text = new_text
            return

        # Фаза 2: cross-run entity — откат к merge
        # В stats уже подсчитана статистика per-run; для full_masked статистику
        # не добавляем повторно, так как это те же данные.
        logger.warning(
            "docx_cross_run_entity_detected",
            note="Paragraph formatting reduced to first-run style due to cross-run sensitive value",
        )
        first_run = paragraph.runs[0]
        first_run.text = full_masked
        for run in paragraph.runs[1:]:
            run.text = ""
