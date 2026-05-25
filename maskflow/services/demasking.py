from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook  # type: ignore[import-untyped]
from openpyxl.worksheet.worksheet import Worksheet  # type: ignore[import-untyped]

from maskflow.core.directory import SUPPORTED_EXTENSIONS
from maskflow.rules.loader import RulesLoader
from maskflow.runtime.paths import resolve_data_path
from maskflow.storage.encrypted_mapping import EncryptedMappingStore
from maskflow.utils.atomic import atomic_write_text
from maskflow.utils.encoding import detect_text_encoding

TEXT_LIKE_EXTENSIONS = {
    ".csv",
    ".json",
    ".log",
    ".sql",
    ".txt",
    ".xml",
}


@dataclass(frozen=True, slots=True)
class DemaskResult:
    replacements: int
    mapping_size: int


class DemaskingService:
    def demask_text(
        self,
        text: str,
        config_path: Path,
    ) -> tuple[str, DemaskResult]:
        mapping = self._load_mapping(config_path)
        demasked, replacements = replace_from_mapping(text, mapping)

        return demasked, DemaskResult(
            replacements=replacements,
            mapping_size=len(mapping),
        )

    def demask_file(
        self,
        source: Path,
        destination: Path,
        config_path: Path,
    ) -> DemaskResult:
        mapping = self._load_mapping(config_path)
        suffix = source.suffix.lower()

        if suffix not in SUPPORTED_EXTENSIONS:
            supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
            raise ValueError(
                f"Unsupported file format: {suffix or '(none)'}. Supported: {supported}"
            )

        destination.parent.mkdir(parents=True, exist_ok=True)

        if suffix in TEXT_LIKE_EXTENSIONS:
            replacements = self._demask_text_file(source, destination, mapping)
        elif suffix == ".docx":
            replacements = self._demask_docx(source, destination, mapping)
        elif suffix == ".xlsx":
            replacements = self._demask_xlsx(source, destination, mapping)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

        return DemaskResult(
            replacements=replacements,
            mapping_size=len(mapping),
        )

    def _load_mapping(self, config_path: Path) -> dict[str, str]:
        config = RulesLoader.load(config_path)

        if not config.reversible_mapping.enabled:
            raise ValueError("reversible_mapping is disabled in config")

        mapping_path = resolve_data_path(config.reversible_mapping.path)
        if not mapping_path.exists():
            raise ValueError(f"Reversible mapping file does not exist: {mapping_path}")

        store = EncryptedMappingStore(
            path=mapping_path,
            encryption_key_env=config.reversible_mapping.encryption_key_env,
        )

        return store.all()

    def _demask_text_file(
        self,
        source: Path,
        destination: Path,
        mapping: dict[str, str],
    ) -> int:
        encoding = detect_text_encoding(source)
        text = source.read_text(encoding=encoding, errors="replace")
        demasked, replacements = replace_from_mapping(text, mapping)
        atomic_write_text(
            destination=destination,
            content=demasked,
            encoding=encoding,
        )

        return replacements

    def _demask_docx(
        self,
        source: Path,
        destination: Path,
        mapping: dict[str, str],
    ) -> int:
        document = Document(str(source))
        replacements = self._demask_docx_document(document, mapping)
        document.save(str(destination))
        return replacements

    def _demask_docx_document(
        self,
        document: DocxDocument,
        mapping: dict[str, str],
    ) -> int:
        replacements = 0

        for paragraph in document.paragraphs:
            replacements += self._demask_docx_paragraph(paragraph, mapping)

        for table in document.tables:
            replacements += self._demask_docx_table(table, mapping)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                replacements += self._demask_docx_paragraph(paragraph, mapping)
            for paragraph in section.footer.paragraphs:
                replacements += self._demask_docx_paragraph(paragraph, mapping)
            for table in section.header.tables:
                replacements += self._demask_docx_table(table, mapping)
            for table in section.footer.tables:
                replacements += self._demask_docx_table(table, mapping)

        return replacements

    def _demask_docx_table(
        self,
        table: Table,
        mapping: dict[str, str],
    ) -> int:
        replacements = 0

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replacements += self._demask_docx_paragraph(paragraph, mapping)
                for nested_table in cell.tables:
                    replacements += self._demask_docx_table(nested_table, mapping)

        return replacements

    def _demask_docx_paragraph(
        self,
        paragraph: Paragraph,
        mapping: dict[str, str],
    ) -> int:
        if not paragraph.runs or not paragraph.text:
            return 0

        full_text, full_replacements = replace_from_mapping(paragraph.text, mapping)
        if full_replacements == 0:
            return 0

        run_results = [replace_from_mapping(run.text, mapping)[0] for run in paragraph.runs]
        per_run_joined = "".join(run_results)

        if per_run_joined == full_text:
            for run, value in zip(paragraph.runs, run_results, strict=True):
                run.text = value
        else:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

        return full_replacements

    def _demask_xlsx(
        self,
        source: Path,
        destination: Path,
        mapping: dict[str, str],
    ) -> int:
        workbook = load_workbook(str(source))
        replacements = 0

        for sheet in workbook.worksheets:
            if not isinstance(sheet, Worksheet):
                continue

            for row in sheet.iter_rows():
                for cell in row:
                    if not isinstance(cell.value, str):
                        continue

                    if cell.value.startswith("="):
                        continue

                    demasked, cell_replacements = replace_from_mapping(cell.value, mapping)
                    if cell_replacements > 0:
                        cell.value = demasked
                        replacements += cell_replacements

        workbook.save(str(destination))
        return replacements


def replace_from_mapping(
    text: str,
    mapping: dict[str, str],
) -> tuple[str, int]:
    result = text
    replacements = 0

    for masked, original in sorted(mapping.items(), key=lambda item: len(item[0]), reverse=True):
        count = result.count(masked)
        if count == 0:
            continue

        result = result.replace(masked, original)
        replacements += count

    return result, replacements
