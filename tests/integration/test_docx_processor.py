from pathlib import Path

from docx import Document

from maskflow.core.engine import MaskingEngine
from maskflow.detectors.email import EmailDetector
from maskflow.formats.docx import DocxProcessor
from maskflow.maskers.hmac_masker import HmacMasker


def test_docx_processor_masks_email(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    destination = tmp_path / "masked.docx"

    document = Document()
    document.add_paragraph("Contact: admin@example.com")
    document.save(str(source))

    engine = MaskingEngine(
        detectors=[EmailDetector()],
        maskers={
            "email": HmacMasker(secret="secret", prefix="EMAIL"),
        },
    )

    processor = DocxProcessor(engine)
    processor.process(str(source), str(destination))

    masked_document = Document(str(destination))
    text = "\n".join(paragraph.text for paragraph in masked_document.paragraphs)

    assert "admin@example.com" not in text
    assert "EMAIL_" in text
